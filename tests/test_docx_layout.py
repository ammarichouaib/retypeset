"""Two-column layout in the Word route.

Both regressions here came from one manuscript, reformatted for IEEE Access,
whose output was unusable in a way that looked like the converter had destroyed
it:

* The **title, authors and abstract appeared in the right-hand column** and the
  introduction continued on the left. The cause was in the source file, not in
  the conversion: Word had written `<w:bidi/>` into the section properties
  because the author has an RTL editing language installed. In one column that
  flag is invisible. Add a second column and it reverses their order.
* The **title block was dragged into the columns** with everything else,
  because the restyler set the column count on the only section there was. The
  previous version knew this and printed instructions telling the author to
  insert the section break by hand in Word -- which is the work they came here
  to avoid.

Neither is about content, and the tests assert that: the paragraph, equation
and image counts must come through untouched.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu, Inches
from lxml import etree

from retypeset.ir import Manuscript, Section, SectionRole
from retypeset.profile import get_profile
from retypeset.render_docx import DocxRestyler, render_docx


def _write(path: Path, paragraphs: list[str], *, bidi: bool = False) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if bidi:
        sectPr = doc.sections[0]._sectPr
        sectPr.append(sectPr.makeelement(qn("w:bidi"), {}))
    doc.save(str(path))
    return path


def _manuscript(front: list[str], body_title: str) -> Manuscript:
    ms = Manuscript()
    ms.meta.title = front[0]
    ms.body = [
        Section(id="s0", level=1, role=SectionRole.ABSTRACT, title_raw="Abstract"),
        Section(id="s1", level=1, role=SectionRole.INTRODUCTION,
                title_raw=body_title),
    ]
    return ms


FRONT = ["A Title That Spans The Page", "A. Author, B. Author",
         "Abstract: something short about the work."]
BODY = ["1 Introduction", "As far back as 1769, the French engineer...",
        "More body text follows here."]


def _sect_columns(path: Path) -> list[int]:
    doc = Document(str(path))
    out = []
    for section in doc.sections:
        cols = section._sectPr.find(qn("w:cols"))
        out.append(int(cols.get(qn("w:num")) or 1) if cols is not None else 1)
    return out


def test_right_to_left_section_layout_is_removed(tmp_path):
    src = _write(tmp_path / "rtl.docx", FRONT + BODY, bidi=True)
    assert Document(str(src)).sections[0]._sectPr.find(qn("w:bidi")) is not None

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    for section in Document(str(out.path)).sections:
        assert section._sectPr.find(qn("w:bidi")) is None
    assert any("right-to-left" in n.lower() for n in out.notes)


def test_title_block_gets_its_own_full_width_section(tmp_path):
    src = _write(tmp_path / "plain.docx", FRONT + BODY)
    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    assert _sect_columns(out.path) == [1, 2]
    assert any("full width" in n for n in out.notes)
    # The break is a property of an existing paragraph, so nothing is inserted
    # into the text.
    assert len(Document(str(out.path)).paragraphs) == len(FRONT + BODY)
    assert not out.unsupported or all(
        "two columns, including the title" not in u for u in out.unsupported)


def test_a_single_column_journal_is_untouched_by_the_split(tmp_path):
    src = _write(tmp_path / "plain.docx", FRONT + BODY)
    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("elsevier_generic"), tmp_path / "out.docx")
    assert _sect_columns(out.path) == [1]


def test_without_a_recognisable_body_start_it_says_so(tmp_path):
    """No guessing: if the boundary is unknown, the author is told what to do."""
    src = _write(tmp_path / "plain.docx", FRONT + ["Some prose, no headings."])
    ms = Manuscript()
    ms.meta.title = FRONT[0]
    ms.body = [Section(id="s0", level=1, role=SectionRole.UNKNOWN, title_raw="")]
    out = render_docx(src, ms, get_profile("ieee_access"), tmp_path / "out.docx")

    assert _sect_columns(out.path) == [2]
    assert any("could not be located" in u for u in out.unsupported)


def test_paragraph_direction_is_kept_where_the_text_is_right_to_left(tmp_path):
    doc = Document()
    doc.add_paragraph(FRONT[0])
    arabic = doc.add_paragraph("ملخص البحث باللغة العربية")
    pPr = arabic._p.get_or_add_pPr()
    pPr.append(pPr.makeelement(qn("w:bidi"), {}))
    english = doc.add_paragraph("An English paragraph pushed RTL by the editor.")
    pPr2 = english._p.get_or_add_pPr()
    pPr2.append(pPr2.makeelement(qn("w:bidi"), {}))
    for text in BODY:
        doc.add_paragraph(text)
    src = tmp_path / "mixed.docx"
    doc.save(str(src))

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    paras = Document(str(out.path)).paragraphs
    by_text = {p.text: p for p in paras}
    kept = by_text["ملخص البحث باللغة العربية"]._p.find(qn("w:pPr")).find(qn("w:bidi"))
    dropped = by_text["An English paragraph pushed RTL by the editor."]._p.find(
        qn("w:pPr")).find(qn("w:bidi"))
    assert kept is not None, "an Arabic paragraph must keep its direction"
    assert dropped is None, "an English paragraph must not be laid out RTL"


# -- oversized figures and hand-padded equation numbers ---------------------
#
# A real manuscript converted through this path (Nebbar V3, hydraulics) came
# back with charts running off the page edge and equation numbers wrapped
# onto their own line, sometimes above the equation. Both traced to content
# sized for the *source* single-column page and carried unchanged into a
# narrower two-column one: a figure inserted at the old full column width,
# and an equation number placed by typing spaces until the cursor reached the
# old margin. These tests reproduce both in a minimal document.


def _png_bytes(width_px: int, height_px: int) -> bytes:
    from PIL import Image
    buf = io.BytesIO()
    Image.new("RGB", (width_px, height_px), "white").save(buf, format="PNG")
    return buf.getvalue()


def _write_with_figure(path: Path, width_emu: int, height_emu: int) -> Path:
    doc = Document()
    for text in FRONT:
        doc.add_paragraph(text)
    for text in BODY:
        doc.add_paragraph(text)
    doc.add_picture(io.BytesIO(_png_bytes(400, 200)),
                     width=Emu(width_emu), height=Emu(height_emu))
    doc.save(str(path))
    return path


def _oMath_with_padding(spaces: int, number: str) -> etree._Element:
    m = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pad = " " * spaces
    xml = (
        f'<m:oMathPara xmlns:m="{m}" xmlns:w="{w}">'
        f'<m:oMath><m:r><m:t>x=y</m:t></m:r>'
        f'<m:r><m:t xml:space="preserve">{pad}</m:t></m:r>'
        f'<m:d><m:e><m:r><m:t>{number}</m:t></m:r></m:e></m:d>'
        f'</m:oMath></m:oMathPara>'
    )
    return etree.fromstring(xml)


def _write_with_equation(path: Path, spaces: int) -> Path:
    doc = Document()
    for text in FRONT:
        doc.add_paragraph(text)
    for text in BODY:
        doc.add_paragraph(text)
    eq_para = doc.add_paragraph()
    eq_para._p.append(_oMath_with_padding(spaces, "3"))
    doc.save(str(path))
    return path


def _column_width_emu(profile_id: str, columns: int) -> int:
    restyler = DocxRestyler.__new__(DocxRestyler)
    restyler.p = get_profile(profile_id)
    return restyler._column_width_emu(columns)


def test_a_figure_too_wide_for_its_new_column_is_shrunk_to_fit(tmp_path):
    # 6 inches: comfortably full-width in the single-column source, and wider
    # than any two-column journal's column.
    original_cx, original_cy = Inches(6), Inches(3)
    src = _write_with_figure(tmp_path / "wide.docx", int(original_cx), int(original_cy))

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    extent = Document(str(out.path)).element.body.find(f".//{qn('wp:extent')}")
    new_cx, new_cy = int(extent.get("cx")), int(extent.get("cy"))
    col_width = _column_width_emu("ieee_access", 2)

    assert new_cx <= col_width * 1.01, "the figure must fit its column"
    # Aspect ratio locked: width and height shrink by the same factor.
    assert abs((new_cx / new_cy) - (original_cx / original_cy)) < 1e-6
    assert any("figure" in n.lower() and "column" in n.lower()
               for n in out.notes)


def test_a_figure_that_already_fits_is_left_alone(tmp_path):
    col_width = _column_width_emu("ieee_access", 2)
    small_cx, small_cy = int(col_width * 0.8), int(col_width * 0.8 * 0.5)
    src = _write_with_figure(tmp_path / "small.docx", small_cx, small_cy)

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    extent = Document(str(out.path)).element.body.find(f".//{qn('wp:extent')}")
    assert int(extent.get("cx")) == small_cx
    assert not any("scaled down" in n for n in out.notes)


def test_hand_typed_equation_number_padding_becomes_a_tab(tmp_path):
    src = _write_with_equation(tmp_path / "eq.docx", spaces=120)

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    doc = Document(str(out.path))
    eq_para = next(p for p in doc.paragraphs if p._p.find(f".//{qn('m:oMath')}") is not None)
    texts = [t.text for t in eq_para._p.findall(f".//{qn('m:t')}")]
    assert "\t" in texts, "the space run must become a real tab character"
    assert not any(t and t.strip() == "" and len(t) > 1 for t in texts), (
        "no long literal space run should remain")

    tabs = eq_para._p.find(qn("w:pPr")).find(qn("w:tabs"))
    tab = tabs.find(qn("w:tab"))
    assert tab.get(qn("w:val")) == "right"
    col_width_twips = _column_width_emu("ieee_access", 2) // 635
    assert int(tab.get(qn("w:pos"))) == col_width_twips
    assert any("equation number" in n.lower() for n in out.notes)


def test_a_short_equation_number_gap_is_left_alone(tmp_path):
    """A handful of spaces is ordinary spacing, not a margin-reaching hack."""
    src = _write_with_equation(tmp_path / "eq.docx", spaces=3)

    out = render_docx(src, _manuscript(FRONT, "1 Introduction"),
                      get_profile("ieee_access"), tmp_path / "out.docx")

    doc = Document(str(out.path))
    eq_para = next(p for p in doc.paragraphs if p._p.find(f".//{qn('m:oMath')}") is not None)
    texts = [t.text for t in eq_para._p.findall(f".//{qn('m:t')}")]
    assert "   " in texts
    assert not any("equation number" in n.lower() for n in out.notes)
