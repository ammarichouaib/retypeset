"""
retypeset.render_docx -- produce a journal-styled .docx.

Design decision: RESTYLE THE ORIGINAL, DO NOT REBUILD IT.

The obvious approach is to walk the IR and emit a fresh document with
python-docx. That is the wrong choice for Word output, for one decisive reason:
python-docx cannot write OMML. Every equation would have to be round-tripped
LaTeX -> MathML -> OMML, which needs Word's own `MML2OMML.XSL`, and any gap in
that chain turns an equation into either a picture or plain text. On a
manuscript with 134 equations that is not a risk worth taking.

Instead we open the author's own .docx and change only presentation: fonts,
sizes, line spacing, margins, column count, line numbering, and the styles
attached to headings and captions. Equations, figures, tables, footnotes and
field codes are never touched, so they cannot be damaged. The IR is still used
-- it tells us which paragraphs are headings, which are captions, and what the
target journal requires -- but it is a map, not a source.

What this cannot do: reorder sections, renumber references into a different
citation style, or synthesise front matter the source lacks. Those need the
rebuild path and are listed in the report rather than attempted silently.
"""

from __future__ import annotations

import copy
import re
import shutil
import stat
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from . import cleanup
from .ir import Manuscript, SectionRole

# Roles that open the body of a paper. The first of these marks the end of the
# title block, which is where a two-column journal wants its section break.
_BODY_ROLES = frozenset({
    SectionRole.INTRODUCTION, SectionRole.RELATED_WORK, SectionRole.THEORY,
    SectionRole.METHODS, SectionRole.EXPERIMENTAL, SectionRole.RESULTS,
    SectionRole.RESULTS_DISCUSSION, SectionRole.DISCUSSION,
    SectionRole.NOMENCLATURE,
})

# "1 Introduction", "I. INTRODUCTION", "1. Background" -- the fallback when no
# role was resolved, which is common in a manuscript with no heading styles.
_OPENS_BODY_RE = re.compile(r"^\s*(?:1|I)\s*[.)\s]\s*\S|^\s*introduction\b",
                            re.I)

# Arabic, Hebrew, Syriac, Thaana and the Arabic presentation forms. A paragraph
# holding any of these keeps its own direction; everything else is being pushed
# right-to-left by an editor setting rather than by its content.
_RTL_RE = re.compile(r"[\u0590-\u08FF\uFB1D-\uFDFF\uFE70-\uFEFF]")


def _has_rtl_text(text: str) -> bool:
    return bool(_RTL_RE.search(text or ""))
from .profile import JournalProfile

_CAPTION_RE = re.compile(
    r"^\s*(fig(?:ure)?|tab(?:le)?|scheme|chart)\s*\.?\s*[A-Z]?\d+", re.I
)


@dataclass
class DocxResult:
    path: Path
    changed_paragraphs: int = 0
    notes: list[str] = field(default_factory=list)
    unsupported: list[str] = field(default_factory=list)


class DocxRestyler:
    def __init__(self, source_docx: str | Path, ms: Manuscript,
                 profile: JournalProfile, *, strip_furniture: bool = True):
        self.src = Path(source_docx)
        self.ms = ms
        self.p = profile
        self.strip_furniture = strip_furniture
        self.notes: list[str] = []
        self.unsupported: list[str] = []
        self._n = 0

    # -- public ------------------------------------------------------------

    def render(self, out_path: str | Path) -> DocxResult:
        out = Path(out_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        if out.exists():
            # A previous run may have left a read-only copy behind.
            out.chmod(out.stat().st_mode | stat.S_IWUSR)
            out.unlink()

        if self.strip_furniture:
            # Remove the source journal's identity before restyling. Correct
            # fonts and margins do not help if the page still carries another
            # journal's logo and running citation line.
            res = cleanup.clean(
                self.src, out, title=self.ms.meta.title,
                first_line_indent_mm=0.0,
            )
            self.notes.extend(res.notes)
            if res.removed_paragraphs:
                self.notes.append(
                    f"{len(res.removed_paragraphs)} boilerplate paragraph(s) "
                    "removed: " + "; ".join(
                        x.split(":")[0] for x in dict.fromkeys(
                            r.split(":")[0] for r in res.removed_paragraphs))
                )
            if res.removed_images:
                self.notes.append(
                    f"{res.removed_images} masthead logo/image(s) removed.")
            if res.reindented:
                self.notes.append(
                    f"{res.reindented} paragraph(s) had inherited left/right "
                    "indents cleared.")
        else:
            shutil.copy2(self.src, out)

        # copy2 preserves permissions, and an uploaded source is often
        # read-only; the copy must be writable for python-docx to save it.
        out.chmod(out.stat().st_mode | stat.S_IWUSR)

        doc = Document(str(out))
        self._page_setup(doc)
        self._base_style(doc)
        self._body(doc)
        self._line_numbers(doc)
        self._report_unsupported()
        doc.save(str(out))
        return DocxResult(out, self._n, self.notes, self.unsupported)

    # -- page --------------------------------------------------------------

    def _page_setup(self, doc: Document) -> None:
        d = self.p.docx
        w, h = (Mm(210), Mm(297)) if d.page_size == "a4" else (Mm(215.9), Mm(279.4))
        for section in doc.sections:
            section.page_width, section.page_height = w, h
            section.top_margin = Mm(d.margins_mm.get("top", 25))
            section.bottom_margin = Mm(d.margins_mm.get("bottom", 25))
            section.left_margin = Mm(d.margins_mm.get("left", 25))
            section.right_margin = Mm(d.margins_mm.get("right", 25))

        self._clear_bidi(doc)
        self._apply_columns(doc, d.columns)
        self.notes.append(
            f"Page set to {d.page_size.upper()}, margins "
            f"{d.margins_mm.get('left', 25):g} mm."
        )

    def _clear_bidi(self, doc: Document) -> None:
        r"""Remove right-to-left layout from the section properties.

        Word writes `<w:bidi/>` into `sectPr` whenever the author has an RTL
        editing language enabled -- Arabic or Hebrew -- even when every word in
        the document is English. In a single-column manuscript it is invisible,
        which is why it survives to submission unnoticed.

        Add a second column and it stops being invisible: the columns fill
        right to left. The title, authors and abstract appear in the *right*
        column and the introduction continues in the left, which reads as a
        catastrophically broken conversion and is in fact one flag in the
        source file. Every journal this tool targets sets text left to right,
        so the flag is cleared with the rest of the page setup.

        Paragraph-level direction is left alone unless the paragraph contains
        no RTL script at all: a genuinely bilingual manuscript -- an Arabic
        abstract, say -- must keep its direction.
        """
        cleared = 0
        for section in doc.sections:
            for tag in ("w:bidi", "w:rtlGutter"):
                el = section._sectPr.find(qn(tag))
                if el is not None:
                    section._sectPr.remove(el)
                    cleared += 1

        paras = 0
        for para in doc.paragraphs:
            pPr = para._p.find(qn("w:pPr"))
            if pPr is None:
                continue
            bidi = pPr.find(qn("w:bidi"))
            if bidi is None or _has_rtl_text(para.text):
                continue
            pPr.remove(bidi)
            paras += 1

        if cleared or paras:
            self.notes.append(
                f"Right-to-left layout removed from {cleared} section "
                f"element(s)" + (f" and {paras} paragraph(s)" if paras else "")
                + ". Word adds this when an RTL editing language is enabled; "
                "in two columns it makes the text flow into the right-hand "
                "column first.")

    def _apply_columns(self, doc: Document, want: int) -> None:
        """Set column counts without destroying an existing title-block layout.

        A two-column journal manuscript is not two columns throughout: the
        title, authors and abstract span the full page, and a continuous section
        break switches the body to two columns. Forcing `w:num` on every section
        collapses that title block into the left column -- which is exactly what
        made the first IEEE output look wrong.

        So: if the document already varies its column count across sections, the
        author has a deliberate layout and we only normalise the multi-column
        sections. Only a uniformly single-column document gets the profile's
        count applied throughout.
        """
        counts = [self._column_count(s) for s in doc.sections]
        varies = len(set(counts)) > 1

        # Collapsing to a single column is always safe: there is no title block
        # to destroy, because one column *is* full width. Only expansion needs
        # the guard. The first version of this rule was symmetric, which left a
        # single-column journal like Elsevier with a two-column manuscript.
        if want == 1:
            for section in doc.sections:
                self._columns(section, 1)
            if any(c > 1 for c in counts):
                self.notes.append(
                    f"Collapsed {sum(1 for c in counts if c > 1)} multi-column "
                    "section(s) to single column, as this journal requires."
                )
            else:
                self.notes.append("Single column throughout.")
            return

        if varies:
            changed = 0
            for section, n in zip(doc.sections, counts):
                if n > 1 and n != want and want > 1:
                    self._columns(section, want)
                    changed += 1
            self.notes.append(
                f"Existing title-block layout preserved: "
                f"{counts.count(1)} full-width section(s), "
                f"{sum(1 for c in counts if c > 1)} multi-column section(s)"
                + (f"; {changed} normalised to {want} columns." if changed else ".")
            )
            return

        if want >= 2 and self._split_front_matter(doc):
            for i, section in enumerate(doc.sections):
                self._columns(section, 1 if i == 0 else want)
            self.notes.append(
                f"Title block kept full width: a continuous section break was "
                f"inserted before the first body section, and the rest of the "
                f"document set to {want} columns.")
            return

        for section in doc.sections:
            self._columns(section, want)
        self.notes.append(f"Set to {want} column(s) throughout.")
        if want >= 2:
            self.unsupported.append(
                "The whole document is now two columns, including the title and "
                "abstract, because the first body section could not be located "
                "automatically. Journals expect the title block full width: in "
                "Word, put the cursor after the abstract and insert "
                "Layout > Breaks > Continuous, then set the first section back "
                "to one column."
            )

    def _first_body_paragraph(self, doc: Document):
        """The paragraph that starts the body, i.e. the end of the title block.

        Taken from the parsed manuscript rather than from the formatting: the
        heading that opens the body is known semantically, and matching it by
        text is more reliable than guessing from style names in a file whose
        styles are the reason we are here.
        """
        titles = [s.title_raw for s in self.ms.body
                  if s.title_raw and s.role in _BODY_ROLES]
        if not titles:
            titles = [s.title_raw for s in self.ms.body
                      if s.title_raw and _OPENS_BODY_RE.match(s.title_raw.strip())]
        wanted = {self._norm(t) for t in titles[:1]}
        if not wanted:
            return None
        for para in doc.paragraphs:
            if self._norm(para.text) in wanted:
                return para
        return None

    def _split_front_matter(self, doc: Document) -> bool:
        r"""Put a continuous section break between the title block and the body.

        This is the step that turns a mechanically correct two-column document
        into one that looks like the journal's own. Without it the title,
        authors and abstract are dragged into the left column and the reader
        meets the paper sideways -- the previous version simply printed
        instructions telling the author to do it by hand in Word, which is the
        work they came here to avoid.

        In OOXML a section break *is* a paragraph: `w:sectPr` inside the
        `w:pPr` of the last paragraph of the section it closes. So the body's
        own properties are copied, set to one column and marked continuous,
        and attached to the paragraph before the first body heading. Nothing is
        added to the text, and the document keeps exactly one visible layout
        change.
        """
        first_body = self._first_body_paragraph(doc)
        if first_body is None:
            return False
        previous = first_body._p.getprevious()
        while previous is not None and previous.tag != qn("w:p"):
            previous = previous.getprevious()
        if previous is None:
            return False
        if len(doc.sections) > 1:
            return False                     # the author already split it

        body_sectPr = doc.sections[-1]._sectPr
        front = copy.deepcopy(body_sectPr)
        for tag in ("w:headerReference", "w:footerReference", "w:lnNumType"):
            for el in front.findall(qn(tag)):
                front.remove(el)
        stype = front.find(qn("w:type"))
        if stype is None:
            stype = front.makeelement(qn("w:type"), {})
            front.insert(0, stype)
        stype.set(qn("w:val"), "continuous")

        pPr = previous.find(qn("w:pPr"))
        if pPr is None:
            pPr = previous.makeelement(qn("w:pPr"), {})
            previous.insert(0, pPr)
        pPr.append(front)
        return True

    @staticmethod
    def _column_count(section) -> int:
        cols = section._sectPr.find(qn("w:cols"))
        if cols is None:
            return 1
        try:
            return int(cols.get(qn("w:num")) or 1)
        except ValueError:
            return 1

    def _columns(self, section, n: int) -> None:
        """Set the column count. python-docx has no API for this.

        Setting `w:num` alone is not enough. When `w:equalWidth="0"` the element
        carries explicit `<w:col>` children giving each column's width, and those
        take precedence: a document rewritten to `w:num="1"` while still holding
        a half-width `<w:col>` renders as two columns. The children must go.
        """
        sectPr = section._sectPr
        cols = sectPr.find(qn("w:cols"))
        if cols is None:
            cols = sectPr.makeelement(qn("w:cols"), {})
            sectPr.append(cols)
        for child in list(cols):
            cols.remove(child)
        cols.set(qn("w:num"), str(max(1, n)))
        cols.set(qn("w:space"), "425")          # ~7.5 mm gutter, in twips
        cols.set(qn("w:equalWidth"), "1")

    def _line_numbers(self, doc: Document) -> None:
        if not self.p.docx.line_numbers:
            return
        for section in doc.sections:
            sectPr = section._sectPr
            ln = sectPr.find(qn("w:lnNumType"))
            if ln is None:
                ln = sectPr.makeelement(qn("w:lnNumType"), {})
                sectPr.append(ln)
            ln.set(qn("w:countBy"), "1")
            ln.set(qn("w:restart"), "continuous")
            ln.set(qn("w:distance"), "360")
        self.notes.append("Continuous line numbering enabled (required for review).")

    # -- styles ------------------------------------------------------------

    def _base_style(self, doc: Document) -> None:
        d = self.p.docx
        normal = doc.styles["Normal"]
        normal.font.name = d.body_font
        normal.font.size = Pt(d.body_size_pt)
        # East-Asian and complex-script font names live in rPr, not the API.
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), d.body_font)

        pf = normal.paragraph_format
        pf.line_spacing = d.line_spacing
        pf.line_spacing_rule = (WD_LINE_SPACING.DOUBLE if d.line_spacing >= 2
                                else WD_LINE_SPACING.MULTIPLE)
        pf.space_after = Pt(0)
        self.notes.append(
            f"Body set to {d.body_font} {d.body_size_pt:g} pt, "
            f"line spacing {d.line_spacing:g}."
        )

    # -- body --------------------------------------------------------------

    def _body(self, doc: Document) -> None:
        d = self.p.docx
        headings = {self._norm(s.title_raw): s
                    for s in self.ms.iter_sections() if s.title_raw}

        for para in doc.paragraphs:
            text = self._norm(para.text)
            if not text:
                continue

            sec = headings.get(text)
            if sec is not None:
                self._apply_heading(para, sec.level)
                self._n += 1
                continue

            if _CAPTION_RE.match(para.text.strip()):
                self._apply_caption(para)
                self._n += 1
                continue

            self._apply_body(para)
            self._n += 1

        # Table text inherits Normal but often carries hard-coded fonts.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = d.body_font
                            if run.font.size and run.font.size.pt > d.body_size_pt:
                                run.font.size = Pt(d.body_size_pt)

    def _apply_body(self, para) -> None:
        d = self.p.docx
        pf = para.paragraph_format
        pf.line_spacing = d.line_spacing
        pf.space_after = Pt(0)
        for run in para.runs:
            # Only override the family and oversized text: bold, italic and
            # sub/superscript carry meaning and must survive.
            run.font.name = d.body_font
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rfonts.set(qn(attr), d.body_font)
            if run.font.size is None or abs(run.font.size.pt - d.body_size_pt) > 0.1:
                run.font.size = Pt(d.body_size_pt)

    def _apply_heading(self, para, level: int) -> None:
        d = self.p.docx
        try:
            para.style = f"Heading {min(max(level, 1), 4)}"
        except KeyError:
            self.unsupported.append(
                f"Heading style level {level} is not defined in this document; "
                "the heading was formatted directly instead."
            )
        sizes = {1: 2.0, 2: 1.0, 3: 0.0, 4: 0.0}
        for run in para.runs:
            run.font.name = d.body_font
            run.font.size = Pt(d.body_size_pt + sizes.get(level, 0.0))
            run.font.bold = True
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(5)

    def _apply_caption(self, para) -> None:
        d = self.p.docx
        for run in para.runs:
            run.font.name = d.body_font
            run.font.size = Pt(max(7.0, d.body_size_pt - 1.0))
        para.paragraph_format.line_spacing = 1.0
        para.paragraph_format.space_after = Pt(6)

    # -- reporting ---------------------------------------------------------

    def _report_unsupported(self) -> None:
        s = self.p.structure
        order = [r for r in s.section_order]
        present = [x.role for x in self.ms.body if x.role in order]
        if present != sorted(present, key=order.index):
            self.unsupported.append(
                "Section order does not match the journal's expected sequence. "
                "Restyling never moves content - reorder the sections in Word."
            )

        missing = [r.value for r in s.required_sections
                   if not any(x.role is r for x in self.ms.iter_sections())
                   and r not in (SectionRole.ABSTRACT, SectionRole.KEYWORDS)]
        if missing:
            self.unsupported.append(
                "Sections this journal requires but the manuscript lacks: "
                + ", ".join(missing) + ". These must be written, not formatted."
            )

        if any(i.code == "manual_citations" for i in self.ms.issues):
            self.unsupported.append(
                f"Citations are plain text and this journal uses "
                f"{self.p.references.style} style. Restyling cannot renumber or "
                "reformat citations that are not reference-manager fields."
            )

        bad = [f.id for f in self.ms.figures
               if f.fmt.lower() in [x.lower() for x in self.p.figures.rejected_formats]]
        if bad:
            self.unsupported.append(
                f"{len(bad)} figure(s) are in a format this journal rejects "
                f"({', '.join(bad[:8])}). Replace them in the output document."
            )

    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", s or "").strip()


def render_docx(source_docx: str | Path, ms: Manuscript, profile: JournalProfile,
                out_path: str | Path, *, strip_furniture: bool = True) -> DocxResult:
    return DocxRestyler(source_docx, ms, profile,
                        strip_furniture=strip_furniture).render(out_path)
